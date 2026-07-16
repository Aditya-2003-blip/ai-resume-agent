import inspect
if not hasattr(inspect, 'getargspec'):
    def getargspec(func):
        full = inspect.getfullargspec(func)
        from collections import namedtuple
        ArgSpec = namedtuple('ArgSpec', ['args', 'varargs', 'keywords', 'defaults'])
        return ArgSpec(full.args, full.varargs, full.varkw, full.defaults)
    inspect.getargspec = getargspec

import warnings
warnings.filterwarnings("ignore", category=FutureWarning)

import streamlit as st
st.config.set_option("server.maxUploadSize", 1024)

from google import genai
import os
import pypdf
import docx2txt
from io import BytesIO
from docx import Document

# Initialize the Gemini client using Google Gen AI SDK
api_key = st.secrets.get("GOOGLE_API_KEY") or os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY")
client = None
if api_key:
    client = genai.Client(api_key=api_key)
else:
    try:
        client = genai.Client()
    except Exception:
        client = None

# Universal text extraction engine for PDFs and Word Docs
def extract_text_from_file(uploaded_file) -> str:
    filename = uploaded_file.name
    try:
        if filename.lower().endswith('.pdf'):
            reader = pypdf.PdfReader(uploaded_file)
            extracted_text = ""
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
            return extracted_text
        elif filename.lower().endswith('.docx'):
            text = docx2txt.process(uploaded_file)
            return text
        else:
            return "Unsupported file format."
    except Exception as e:
        return f"Error reading file {filename}: {str(e)}"

# Stable Word Document Generator
def create_docx_report(text_content: str) -> BytesIO:
    doc = Document()
    doc.add_heading("ATS Evaluation & Recruitment Analytics Report", level=0)
    for line in text_content.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('* '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- STREAMLIT WORKSPACE ARCHITECTURE ---
st.set_page_config(page_title="Enterprise AI ATS Dashboard", page_icon="📈", layout="wide")

st.title("📈 Enterprise AI Resume Screening & Optimization Dashboard")
st.write("Upload candidate profiles to rank match compatibility, pinpoint technical gaps, and receive automated optimization suggestions.")

# Sidebar - Clean Candidate Profiles Area
st.sidebar.header("📁 Candidate Profiles")
uploaded_files = st.sidebar.file_uploader(
    "Upload Resumes (PDF or Word format, up to 10 files):", 
    type=["pdf", "docx"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.sidebar.success(f"Successfully staged {len(uploaded_files)} profile(s).")
else:
    st.sidebar.info("Staging Area: Awaiting documents...")

# Main Panel Configuration
st.subheader("📋 Target Job Specification")
job_description_input = st.text_area(
    "Paste the enterprise role requirements and technical qualifications below:", 
    height=180, 
    placeholder="Looking for an engineer with expertise in..."
)

if st.button("Check 🕵️‍♂️"):
    if not client:
        st.error("⚠️ Configuration missing: GOOGLE_API_KEY or GEMINI_API_KEY is not set.")
    elif not uploaded_files:
        st.warning("⚠️ Action required: Please upload at least one resume profile in the sidebar.")
    elif not job_description_input.strip():
        st.warning("⚠️ Action required: Please supply a Target Job Specification.")
    else:
        file_count = len(uploaded_files)
        
        with st.spinner("🕵️‍♂️ Processing profiles through screening intelligence..."):
            resume_manifests = []
            for uploaded_file in uploaded_files:
                extracted_text = extract_text_from_file(uploaded_file)
                if not extracted_text.startswith("Error"):
                    resume_manifests.append(f"--- CANDIDATE FILE: {uploaded_file.name} ---\n{extracted_text}\n")
            
            all_resumes_combined_text = "\n".join(resume_manifests)
            
            # Dynamically alter prompting based on single vs multiple candidates
            if file_count == 1:
                verdict_instruction = (
                    "Provide a top section titled '## 🏆 Executive Candidate Summary'. "
                    "Since there is only one candidate, do NOT perform a winner selection or comparison. "
                    "Instead, provide a concise 3-sentence summary highlighting their high-level core suitability and readiness for this specific position."
                )
            else:
                verdict_instruction = (
                    "Provide a top section titled '## 🏆 Final Recruitment Verdict & Recommendation'. "
                    "Since there are multiple candidates, look at them collectively, provide a definitive non-diplomatic recommendation ranking them, "
                    "and select a clear winner candidate with a short logical justification."
                )
                
            master_prompt = (
                f"You are an elite corporate technical recruiter and strict ATS evaluation engine. "
                f"You must evaluate all submitted resumes against the Target Job Description.\n\n"
                f"CRITICAL FOR INDIVIDUAL PROFILE FORMAT: For every single candidate analyzed, you MUST use the exact same Markdown structural headers. "
                f"Each individual analysis must strictly contain these exact sections:\n"
                f"### 📊 ATS Match Analysis\n"
                f"* **Candidate Name**: [Name]\n"
                f"* **Target Position**: [Job Title]\n"
                f"* **Estimated Match Score**: [Score/100]\n"
                f"### 🔍 Detailed Evaluation\n"
                f"Provide 3 specific strengths and 3 clear missing gaps.\n"
                f"### 🗣️ Tailored Interview Questions\n"
                f"Provide 3 precise questions testing the uncovered gaps.\n"
                f"### 🛠️ Actionable Resume Revisions\n"
                f"Provide 3-4 highly specific bullet points outlining keywords or phrasing changes to improve their score.\n\n"
                f"TASK:\n"
                f"1. Read all candidate resumes below.\n"
                f"2. {verdict_instruction}\n"
                f"3. Provide individual reports for each candidate using the exact standard structural headers requested.\n\n"
                f"JOB DESCRIPTION:\n{job_description_input}\n\n"
                f"CANDIDATE RESUMES TO PROCESS:\n{all_resumes_combined_text}"
            )
            
            # Unified direct call using the new Google Gen AI SDK
            response = client.models.generate_content(
                model="gemini-1.5-flash",
                contents=master_prompt
            )
            final_compiled_report = response.text
            
            # Metrics parsing logic for the scoreboard
            scoreboard_data = []
            for uploaded_file in uploaded_files:
                filename = uploaded_file.name
                score = "Evaluated"
                for line in final_compiled_report.split("\n"):
                    if "Estimated Match Score:" in line and filename.split(".")[0][:5].lower() in line.lower():
                        potential_score = line.split(":")[-1].strip().replace("**", "")
                        if "/" in potential_score:
                            score = potential_score
                scoreboard_data.append({"Candidate File": filename, "ATS Match Score": score})

        # --- OUTPUT INTERFACE DESIGN LAYOUT ---
        st.success("✅ Analysis Complete!")
        
        st.subheader("📊 Candidate Ranking Scoreboard")
        st.table(scoreboard_data)
        
        st.subheader("💾 Export Assessment Intelligence")
        st.download_button(
            label="📥 Download Full Report as Word (.docx)",
            data=create_docx_report(final_compiled_report),
            file_name="ATS_Recruitment_Comparative_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        st.subheader("🏆 Executive Selection & Strategic Analytics")
        st.markdown(final_compiled_report)
